from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# name = 'Dang Yu'
# phone_number = '15998218724'
# email = 'dangyu@nankai.edu.cn'
# profile picture
document.add_picture(
    'profile.jpg',
    width=Inches(2.0))
# name, phone number and email
name = input('What is your name?')
speak('Hello' + name + 'How are you today?')
phone_number = input('What is your phone number')
speak('Thanks' + ' So your phone number is ' + phone_number)
email = input('What is your email')
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About Me')
document.add_paragraph(input(
    'Tell me about yourself?'
))
# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()
company = input('Company name')
from_date = input('From date ')
to_date = input('To date ')
p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
experience_details: str = input(
    'describe your experience at ' + company
)
p.add_run(experience_details)

# more experience
while True:
    has_more_experience = input(
        'Do you have more experience? Yes or No :')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Company name')
        from_date = input('From date ')
        to_date = input('To date ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        experience_details: str = input(
            'describe your experience at ' + company
        )
        p.add_run(experience_details)
    else:
        break  # break out of the while loop
# add skills
document.add_heading('Skills')
p = document.add_paragraph()
skill = input('Your skill ')
p.add_run(skill)
p.style = 'List Bullet'
while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        p = document.add_paragraph()
        skill = input('Your skill ')
        p.add_run(skill)
        p.style = 'List Bullet'
    else:
        break

    # footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigos code'

document.save('cv.docx')
